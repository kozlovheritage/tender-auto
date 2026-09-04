"""Извлечение текста: docx/doc/pdf/xlsx/rtf/odt/txt + лимиты + критичность.

Headless: без print(), только logging. Возвращает данные, а не печатает их.
Заменяет монолитные extract_text_from_file / extract_texts_from_paths.
"""
from __future__ import annotations
import os
import re
import zipfile
from pathlib import Path
from typing import Iterable, List, Tuple

from tendercore.log import get_logger

log = get_logger("extract")

# ── Лимиты (калибровка из монолита) ──
# ── Этап 2.2: лимиты из config/settings.toml (паритет с монолитом) ──
try:
    from tendercore.config import get as _cfg_get_extr
    MAX_TEXT_TOTAL = int(_cfg_get_extr('extraction', 'max_text_total', default=200000) or 200000)
    MAX_TEXT_PER_FILE = int(_cfg_get_extr('extraction', 'max_text_per_file', default=40000) or 40000)
except Exception:
    MAX_TEXT_TOTAL = 200000
    MAX_TEXT_PER_FILE = 40000

OCR_MAX_PAGES = 15

CRITICAL_KEYWORDS = ["техническое задание", "тз", "спецификация",
                     "проект договора", "договор", "документация",
                     "описание объекта закупки", "обоснование нмц"]

_ERR_PREFIXES = ("[Ошибка", "[OCR не дал текста]", "[OCR недоступен]",
                 "[PDF parsing not available]", "[Ошибка чтения .doc]")

# ── Опциональные зависимости (деградация без падения) ──
try:
    import fitz
    _FITZ = True
except ImportError:
    _FITZ = False
try:
    from striprtf.striprtf import rtf_to_text
    _RTF = True
except ImportError:
    _RTF = False
try:
    import pytesseract
    from PIL import Image
    _OCR = True
except ImportError:
    _OCR = False


# ── Чистые функции (тестируемо без внешних зависимостей) ──

def apply_limits(chunks: List[Tuple[str, str]],
                 max_per_file: int = MAX_TEXT_PER_FILE,
                 max_total: int = MAX_TEXT_TOTAL
                 ) -> Tuple[str, List[Tuple[str, int, int]]]:
    """Обрезка по лимитам. chunks: [(filename, text)].

    Возвращает (combined_text, stats), stats: [(filename, kept, total)].
    """
    out, stats = [], []
    total = 0
    for name, text in chunks:
        text = text or ""
        kept = text[:max_per_file]
        room = max_total - total
        if room <= 0:
            stats.append((name, 0, len(text)))
            continue
        kept = kept[:room]
        total += len(kept)
        stats.append((name, len(kept), len(text)))
        out.append(kept)
    return "\n".join(out), stats


def _norm_name(name: str) -> str:
    """Нормализация имени файла: нижний регистр, _ - . → пробел."""
    return name.lower().replace("_", " ").replace("-", " ").replace(".", " ")


def is_critical_file(failed_name: str, success_files) -> bool:
    """Файл критичен, если сам относится к ТЗ/договору и не покрыт другим успешным."""
    fn = _norm_name(failed_name)
    if not any(kw in fn for kw in CRITICAL_KEYWORDS):
        return False
    for other in success_files:
        if any(kw in _norm_name(other) for kw in CRITICAL_KEYWORDS):
            return False
    return True


# ── Экстракторы по форматам ──

def extract_docx(path) -> str:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = " | ".join(c.text.strip() for c in row.cells)
            if cells.strip(" |"):
                parts.append(cells)
    return "\n".join(parts)


def extract_xlsx(path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(path), data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            vals = [str(c) for c in row if c is not None]
            if vals:
                parts.append(" | ".join(vals))
    wb.close()
    return "\n".join(parts)


def extract_rtf(path) -> str:
    if not _RTF:
        return "[Ошибка чтения .rtf: striprtf не установлен]"
    raw = Path(path).read_text(encoding="utf-8", errors="ignore")
    return rtf_to_text(raw)


def extract_odt(path) -> str:
    with zipfile.ZipFile(path) as z:
        xml = z.read("content.xml").decode("utf-8", errors="ignore")
    text = re.sub(r"<[^>]+>", " ", xml)
    return re.sub(r"\s+", " ", text).strip()


def extract_pdf(path) -> str:
    if not _FITZ:
        return "[PDF parsing not available]"
    doc = fitz.open(str(path))
    text = ""
    ocr_pages = []
    for page in doc:
        t = page.get_text()
        if t and len(t.strip()) >= 40:
            text += t
        else:
            ocr_pages.append(page)
    if ocr_pages and _OCR:
        skipped = max(0, len(ocr_pages) - OCR_MAX_PAGES)
        for page in ocr_pages[:OCR_MAX_PAGES]:
            try:
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                t = pytesseract.image_to_string(img, lang="rus+eng")
                if t:
                    text += t + "\n"
            except Exception:
                continue
        if skipped:
            log.warning(f"OCR: {Path(path).name} — пропущено {skipped} стр. (лимит)")
    doc.close()
    return text


def extract_doc(path) -> str:
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(Path(path).resolve()))
        text = doc.Content.Text
        doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()
        return text
    except Exception as e:
        return f"[Ошибка чтения .doc: {e}]"


def extract_text_from_file(path) -> str:
    ext = Path(path).suffix.lower()
    try:
        if ext == ".docx":
            return extract_docx(path)
        if ext == ".doc":
            return extract_doc(path)
        if ext == ".pdf":
            return extract_pdf(path)
        if ext in (".xlsx", ".xls"):
            return extract_xlsx(path)
        if ext == ".rtf":
            return extract_rtf(path)
        if ext == ".odt":
            return extract_odt(path)
        if ext == ".txt":
            return Path(path).read_text(encoding="utf-8", errors="ignore")
        return ""
    except Exception as e:
        return f"[Ошибка чтения {ext or path}: {e}]"


# ── Агрегатор ──

def extract_texts_from_paths(paths, max_total=MAX_TEXT_TOTAL,
                             max_per_file=MAX_TEXT_PER_FILE
                             ) -> Tuple[str, List[str], List[str], List[Tuple[str, str]]]:
    """Возвращает (doc_text, critical_errors, success_files, extraction_errors)."""
    chunks, errors, success = [], [], []
    for p in paths:
        name = os.path.basename(p)
        text = extract_text_from_file(p)
        if text and not text.startswith(_ERR_PREFIXES):
            success.append(name)
            chunks.append((name, text))
            log.info(f"📄 Извлечено {min(len(text), max_per_file)} "
                     f"из {len(text)} символов из {name}")
        else:
            errors.append((name, (text or "пусто")[:100]))
            log.warning(f"⚠️ Ошибка извлечения из {name}: "
                        f"{(text or 'пусто')[:100]}")
    combined, _stats = apply_limits(chunks, max_per_file, max_total)
    critical = [n for n, _ in errors if is_critical_file(n, success)]
    return combined, critical, success, errors