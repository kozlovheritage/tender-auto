"""tendercore.report.full — полное формирование Word-отчёта (зеркало монолита).

Портировано: extract_markdown_table, add_*_table, add_structured_report,
очистка служебных строк, create_tender_report. Headless: только logging.
"""
from __future__ import annotations
import os
import re

from docx import Document
from docx.shared import Pt

from tendercore.log import get_logger

log = get_logger("report")


def extract_markdown_table(text: str):
    """Извлекает строки Markdown-таблицы → list[list[str]] или None."""
    lines = [l.rstrip() for l in text.split('\n')]
    table_lines = [l for l in lines if '|' in l]
    if not table_lines:
        return None
    rows = []
    for line in table_lines:
        if re.match(r'^\s*\|[-:\s|]+\|\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if any(c for c in cells):
            rows.append(cells)
    return rows if len(rows) >= 2 else None


def add_key_value_table(doc, content: str):
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
        if ':' in line:
            key, _, val = line.partition(':')
            p = doc.add_paragraph()
            p.add_run(key.strip() + ': ').bold = True
            p.add_run(val.strip())
        else:
            doc.add_paragraph(line)


def add_single_column_table(doc, content: str):
    for line in content.split('\n'):
        line = line.strip()
        if line and not line.startswith('---'):
            doc.add_paragraph(line, style='List Bullet')


def add_structured_report(doc, response_text: str):
    """Разбирает ответ по «--- Название ---» и добавляет секции в doc."""
    parts = re.split(r'\n(---\s*.+?\s*---)\n', response_text)
    if len(parts) <= 1:
        doc.add_paragraph(response_text)
        return
    preamble = parts[0].strip()
    if preamble:
        for line in preamble.split('\n'):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            if line.startswith('Решение:'):
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
            else:
                p.add_run(line)
    i = 1
    while i < len(parts) - 1:
        section_title = parts[i].strip().strip('-').strip()
        section_content = parts[i + 1].strip() if i + 1 < len(parts) else ""
        i += 2
        doc.add_heading(section_title, level=2)
        if not section_content:
            continue
        table_rows = extract_markdown_table(section_content)
        if table_rows:
            num_cols = max(len(row) for row in table_rows)
            if num_cols > 0:
                t = doc.add_table(rows=0, cols=num_cols)
                t.style = 'Table Grid'
                for row_idx, row_cells in enumerate(table_rows):
                    row_obj = t.add_row()
                    for col_idx, cell_text in enumerate(row_cells):
                        if col_idx < num_cols:
                            cell = row_obj.cells[col_idx]
                            cell.text = cell_text
                            if row_idx == 0:
                                for par in cell.paragraphs:
                                    for run in par.runs:
                                        run.bold = True
        elif ':' in section_content and '\n' in section_content:
            add_key_value_table(doc, section_content)
        else:
            add_single_column_table(doc, section_content)


_SERVICE_LINE_PREFIXES = (
    '=== результаты поиска', '=== конец результатов',
    '=== верифицированные email', '=== используй эти email',
    '--- аналоги бренда', '--- производители глобально',
    '--- азия', '--- азия/турция', '--- европа',
    '--- дистрибьюторы', '--- производитель', '--- производители',
)


def _is_service_line(line: str) -> bool:
    low = line.strip().lower()
    return any(low.startswith(p) for p in _SERVICE_LINE_PREFIXES)


def _clean_suppliers_for_docx(suppliers_text: str) -> list:
    """Убирает служебные строки, сырые URL и HTML-артефакты."""
    if not suppliers_text:
        return []
    clean = []
    for raw in suppliers_text.split('\n'):
        line = raw.strip()
        if not line:
            clean.append('')
            continue
        if _is_service_line(line):
            continue
        if re.match(r'^https?://', line):
            continue
        line = re.sub(r'^(%[0-9A-Fa-f]{2})+', '', line).strip()
        if line.lower().startswith('u003e'):
            line = line[5:].strip()
        if line:
            clean.append(line)
    result, prev_empty = [], False
    for ln in clean:
        if ln == '':
            if not prev_empty:
                result.append(ln)
            prev_empty = True
        else:
            result.append(ln)
            prev_empty = False
    return result


def create_tender_report(reg_number: str, deadline_str: str, response: str,
                         suppliers_text: str, output_dir: str) -> str:
    """Создаёт Word-отчёт; возвращает путь к файлу."""
    doc = Document()
    doc.add_heading(f"Аналитический отчёт по закупке № {reg_number}", level=1)
    doc.add_paragraph(f"Дата окончания подачи заявок: {deadline_str}")
    add_structured_report(doc, response)
    doc.add_heading("Возможные поставщики и контакты", level=2)
    if suppliers_text:
        clean_lines = _clean_suppliers_for_docx(suppliers_text)
        non_empty = [l for l in clean_lines if l.strip()]
        if non_empty:
            for line in clean_lines:
                if line.strip():
                    doc.add_paragraph(line)
        else:
            doc.add_paragraph("(Заполняется вручную)")
    else:
        doc.add_paragraph("(Заполняется вручную)")
    os.makedirs(output_dir, exist_ok=True)
    doc_path = os.path.join(output_dir, f"{reg_number}.docx")
    doc.save(doc_path)
    return doc_path