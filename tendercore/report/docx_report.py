"""Сборка Word-отчёта из структурированного ответа ЛЛМ."""
from __future__ import annotations
import re
from pathlib import Path

from tendercore.log import get_logger

log = get_logger("report")

try:
    from docx import Document
    _DOCX = True
except ImportError:
    _DOCX = False

SECTION_RE = re.compile(r'^---\s*(.+?)\s*---\s*$', re.MULTILINE)


def parse_response_sections(response: str) -> dict:
    """Режет ответ ЛЛМ по маркерам «--- Заголовок ---» → {заголовок: контент}.

    Текст до первого маркера сохраняется под ключом '_preamble' (строка «Решение:»).
    """
    sections = {}
    matches = list(SECTION_RE.finditer(response or ""))
    preamble = (response[:matches[0].start()] if matches else (response or "")).strip()
    if preamble:
        sections["_preamble"] = preamble
    for i, m in enumerate(matches):
        end = matches[i + 1].start() if i + 1 < len(matches) else len(response)
        sections[m.group(1).strip()] = response[m.end():end].strip()
    return sections


def parse_markdown_table(text: str) -> list:
    """Markdown-таблица → список строк-ячеек (разделительная строка отбрасывается)."""
    rows = []
    for line in (text or "").splitlines():
        if "|" not in line:
            continue
        if re.match(r'^\s*\|?\s*[-:|\s]+\|?\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if any(cells):
            rows.append(cells)
    return rows


def build_report(reg_number: str, deadline: str, response: str,
                 suppliers_text: str, path) -> Path:
    """Собирает Word-отчёт; возвращает путь к сохранённому файлу."""
    if not _DOCX:
        raise RuntimeError("python-docx не установлен: pip install python-docx")
    doc = Document()
    doc.add_heading(f"Аналитический отчёт по закупке № {reg_number}", level=1)
    doc.add_paragraph(f"Дата окончания подачи заявок: {deadline}")

    sections = parse_response_sections(response)
    if sections.get("_preamble"):
        doc.add_paragraph(sections["_preamble"])

    for title, content in sections.items():
        if title == "_preamble":
            continue
        doc.add_heading(title, level=2)
        if "|" in content:
            rows = parse_markdown_table(content)
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    for j in range(ncols):
                        table.cell(i, j).text = row[j] if j < len(row) else ""
                continue
        for para in content.splitlines():
            if para.strip():
                doc.add_paragraph(para)

    if suppliers_text and suppliers_text.strip():
        doc.add_heading("Возможные поставщики и контакты", level=2)
        for para in suppliers_text.splitlines():
            if para.strip():
                doc.add_paragraph(para)

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path