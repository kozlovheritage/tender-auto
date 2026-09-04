from tendercore.report.docx_report import (
    build_report, parse_markdown_table, parse_response_sections,
)

RESPONSE = """Решение: участвуем
--- Основная информация ---
Заказчик: Тестовый заказчик
Срок поставки: 45 дней
--- Что нужно поставить ---
| № | Наименование | Артикул | Ед. изм. | Кол-во |
| --- | --- | --- | --- | --- |
| 1 | Снегоочиститель | SR-2500 | шт | 1 |
--- Производитель / бренд / страна ---
UM-Truck
--- Возможные поставщики ---
(заполняется вручную)
"""


def test_parse_sections():
    s = parse_response_sections(RESPONSE)
    assert "Основная информация" in s
    assert "Что нужно поставить" in s
    assert "участвуем" in s["_preamble"]
    assert "Заказчик: Тестовый заказчик" in s["Основная информация"]


def test_parse_markdown_table():
    rows = parse_markdown_table("| № | Название |\n| --- | --- |\n| 1 | Товар |")
    assert rows == [["№", "Название"], ["1", "Товар"]]


def test_build_report(tmp_path):
    out = tmp_path / "32616335374.docx"
    build_report("32616335374", "07.09.2026", RESPONSE,
                 "Поставщик 1 — UM-Truck\nE-mail для поставщика 1 — x@y.ru", out)
    assert out.exists()
    from docx import Document
    d = Document(str(out))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "Аналитический отчёт по закупке № 32616335374" in text
    assert "Дата окончания подачи заявок: 07.09.2026" in text
    assert "Поставщик 1 — UM-Truck" in text
    assert len(d.tables) == 1
    assert d.tables[0].cell(1, 1).text == "Снегоочиститель"