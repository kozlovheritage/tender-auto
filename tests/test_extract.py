import docx
import openpyxl

from tendercore.extract.text import (
    apply_limits, is_critical_file, extract_docx, extract_xlsx,
    extract_text_from_file, extract_texts_from_paths,
)


def test_apply_limits_per_file_and_total():
    chunks = [("a.txt", "x" * 100), ("b.txt", "y" * 100)]
    combined, stats = apply_limits(chunks, max_per_file=60, max_total=100)
    assert stats == [("a.txt", 60, 100), ("b.txt", 40, 100)]
    assert "x" * 60 in combined and "y" * 40 in combined


def test_apply_limits_stops_after_total():
    chunks = [("a", "a" * 10), ("b", "b" * 10)]
    _, stats = apply_limits(chunks, max_per_file=10, max_total=10)
    assert stats[1][1] == 0


def test_is_critical_file():
    assert is_critical_file("Техническое_задание.docx", ["Извещение.docx"])
    assert not is_critical_file("Техническое_задание.docx",
                                ["Техническое_задание_2.docx"])
    assert not is_critical_file("Анкета_участника.docx", [])


def test_docx_roundtrip(tmp_path):
    p = tmp_path / "t.docx"
    d = docx.Document()
    d.add_paragraph("Насос Grundfos CR 32")
    d.save(str(p))
    assert "Grundfos" in extract_docx(p)


def test_xlsx_roundtrip(tmp_path):
    p = tmp_path / "t.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Наименование", "Кол-во"])
    ws.append(["Клапан", 5])
    wb.save(str(p))
    assert "Клапан" in extract_xlsx(p)


def test_rtf_via_dispatcher(tmp_path):
    p = tmp_path / "t.rtf"
    p.write_text(r"{\rtf1\ansi Привет мир\par}", encoding="utf-8")
    assert "Привет" in extract_text_from_file(p)


def test_aggregator_critical(tmp_path):
    good = tmp_path / "Техническое_задание.docx"
    d = docx.Document()
    d.add_paragraph("Спецификация: клапан Danfoss")
    d.save(str(good))
    text, critical, success, errors = extract_texts_from_paths([str(good)])
    assert "Danfoss" in text
    assert critical == [] and errors == []