from docx import Document

from tendercore.report.full import (
    _clean_suppliers_for_docx, _is_service_line, create_tender_report,
    extract_markdown_table,
)

RESP = """Решение: участвуем
--- Основная информация ---
Заказчик: АО «МАШ»
Предмет закупки: Поставка модулей полампового контроля
Количество позиций: 2
НМЦК/НМЦД: 3 797 902,70 рубля с учётом НДС
--- Что нужно поставить ---
| № | Наименование | Артикул/Каталожный номер | Ед. изм. | Кол-во |
| --- | --- | --- | --- | --- |
| 1 | Модуль LMC | LMC-01.1 | шт | 5 |
| 2 | Модуль LLC | LLC-01.1 | шт | 13 |
--- Производитель / бренд / страна ---
Transcon (Чехия)
--- Критичные требования ТЗ ---
- Гарантийный срок не менее 1 года
--- Возможные поставщики ---
(заполняется вручную)
--- Доступность закупки в Китае: Да ---
"""

SUP = """Поставщик 1 — Transcon Electronic Systems, spol. s r.o. (Чехия)
E-mail для поставщика 1 — info@transcon.cz
=== результаты поиска ===
https://example.com/x
"""


def test_extract_markdown_table():
    rows = extract_markdown_table(
        "| № | Наименование | Артикул | Ед. | Кол-во |\n"
        "| --- | --- | --- | --- | --- |\n"
        "| 1 | Модуль LMC | LMC-01.1 | шт | 5 |\n"
        "| 2 | Модуль LLC | LLC-01.1 | шт | 13 |")
    assert rows and len(rows) == 3
    assert rows[1][1] == "Модуль LMC"


def test_extract_markdown_table_none():
    assert extract_markdown_table("просто текст без таблицы") is None


def test_is_service_line():
    assert _is_service_line("=== результаты поиска ===")
    assert _is_service_line("--- производители глобально")
    assert not _is_service_line("Поставщик 1 — Transcon (Чехия)")


def test_clean_suppliers():
    joined = "\n".join(_clean_suppliers_for_docx(SUP))
    assert "===" not in joined and "https://" not in joined
    assert "Поставщик 1 — Transcon" in joined


def test_create_report_structure(tmp_path):
    path = create_tender_report("32616322289", "01.09.2026", RESP, SUP, str(tmp_path))
    doc = Document(path)
    assert any("32616322289" in p.text for p in doc.paragraphs[:2])
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert len(t.rows) == 3 and len(t.columns) == 5
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "Поставщик 1 — Transcon" in texts
    assert "https://" not in texts